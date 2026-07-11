import hashlib
import io
import json
import logging
import re
import secrets
import shutil
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timedelta
from pathlib import Path

import fitz  # PyMuPDF — skan edilmiş PDF səhifələrini şəklə çevirmək üçün
import pytesseract
from docx import Document as DocxDocument
from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from PIL import Image
from pypdf import PdfReader
from sqlalchemy import delete, or_, select
from sqlalchemy.orm import Session, joinedload

from sqlalchemy import text as sql_text

from .ai import analyze_document, answer_question
from .config import settings
from .database import Base, SessionLocal, engine, get_db
from .models import Analysis, Client, CourtEvent, Document, DocumentChunk, User
from .rag import chunk_text, embed_document_text, embed_query, rerank_and_answer
from .schemas import (
    AskRequest,
    CourtEventCreate,
    DeadlineOut,
    DocumentOut,
    LoginRequest,
    RegisterRequest,
)


@asynccontextmanager
async def lifespan(app: FastAPI):
    Path(settings.upload_dir).mkdir(exist_ok=True)
    with engine.begin() as conn:
        # Semantik axtarış üçün pgvector genişlənməsi (postgres image-i pgvector/pgvector:pg16 olmalıdır).
        conn.execute(sql_text("CREATE EXTENSION IF NOT EXISTS vector"))
    Base.metadata.create_all(bind=engine)
    yield


logging.basicConfig(level=logging.INFO)

app = FastAPI(title="LegalAI", version="1.0.0", lifespan=lifespan)
app.mount("/ui", StaticFiles(directory="app/static", html=True), name="ui")

logger = logging.getLogger("legalai")


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    # Əsl xətanı konteyner loglarına yazırıq ki, "docker compose logs api" ilə görünsün —
    # istifadəçiyə isə ümumi mesaj qaytarırıq.
    logger.exception("Gözlənilməz xəta: %s %s", request.method, request.url.path)
    return JSONResponse(status_code=500, content={"detail": "Serverdə gözlənilməz xəta baş verdi. Bir azdan yenidən cəhd edin."})


def _ocr_image(image: Image.Image) -> str:
    return pytesseract.image_to_string(image, lang="aze+eng")


def _ocr_pdf(file_path: Path) -> str:
    """Skan edilmiş (mətni çıxmayan) PDF-lər üçün: hər səhifəni şəklə çevirib OCR edir."""
    pages_text = []
    with fitz.open(str(file_path)) as pdf:
        for page in pdf:
            pixmap = page.get_pixmap(dpi=200)
            image = Image.open(io.BytesIO(pixmap.tobytes("png")))
            pages_text.append(_ocr_image(image))
    return "\n".join(pages_text)


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".txt":
        raw = file_path.read_text(errors="ignore")
    elif suffix in {".jpg", ".jpeg", ".png"}:
        raw = _ocr_image(Image.open(file_path))
    elif suffix == ".pdf":
        raw = "\n".join(page.extract_text() or "" for page in PdfReader(str(file_path)).pages)
        if not raw.strip():
            raw = _ocr_pdf(file_path)
    elif suffix == ".docx":
        raw = "\n".join(p.text for p in DocxDocument(str(file_path)).paragraphs)
    else:
        raise HTTPException(400, "Yalnız PDF, DOCX, TXT və ya şəkil (JPG/PNG) faylları qəbul edilir.")
    # PDF-də cümlələr sətir arası qırılır; boşluq/sətir keçidlərini tək boşluğa endiririk ki,
    # istifadəçi axtarışda yazdığı cümlə mətndə fərqli sətirlərə bölünsə də tapılsın.
    return re.sub(r"\s+", " ", raw).strip()


def serialize(document: Document) -> dict:
    analysis = document.analysis
    if analysis is None:
        analysis_data = {
            "summary": "Bu sənəd üçün analiz mövcud deyil.",
            "risks": [],
            "grammar_issues": [],
            "extracted_dates": [],
            "risk_score": 0,
        }
    else:
        analysis_data = {
            "summary": analysis.summary,
            "risks": json.loads(analysis.risks),
            "grammar_issues": json.loads(analysis.grammar_issues),
            "extracted_dates": json.loads(analysis.extracted_dates),
            "risk_score": analysis.risk_score,
        }
    return {
        "id": document.id,
        "original_filename": document.original_filename,
        "document_type": document.document_type,
        "client_name": document.client.full_name,
        "created_at": document.created_at,
        "analysis": analysis_data,
    }


@app.get("/health")
def health_check():
    return {"status": "ok", "message": "LegalAI server işləyir"}


def _hash_password(password: str, salt: str) -> str:
    return hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 100_000).hex()


def _verify_user_credentials(db: Session, username: str, password: str) -> bool:
    user = db.execute(select(User).where(User.username == username)).scalar_one_or_none()
    if not user:
        return False
    return _hash_password(password, user.salt) == user.password_hash


@app.post("/auth/login")
def login(payload: LoginRequest, db: Session = Depends(get_db)):
    # Sadə giriş: ya .env-dəki tək-istifadəçi (APP_USERNAME/APP_PASSWORD), ya da qeydiyyatdan
    # keçmiş istifadəçilər cədvəli. Bu, real sessiya/token təhlükəsizliyi deyil — yalnız şəxsi
    # kabinetin qarşısında sadə qapıdır.
    if payload.username == settings.app_username and payload.password == settings.app_password:
        return {"ok": True}
    if _verify_user_credentials(db, payload.username.strip(), payload.password):
        return {"ok": True}
    raise HTTPException(401, "İstifadəçi adı və ya parol yanlışdır.")


@app.post("/auth/register")
def register(payload: RegisterRequest, db: Session = Depends(get_db)):
    username = payload.username.strip()
    if username == settings.app_username:
        raise HTTPException(400, "Bu istifadəçi adı artıq mövcuddur.")
    existing = db.execute(select(User).where(User.username == username)).scalar_one_or_none()
    if existing:
        raise HTTPException(400, "Bu istifadəçi adı artıq mövcuddur.")
    salt = secrets.token_hex(16)
    user = User(username=username, salt=salt, password_hash=_hash_password(payload.password, salt))
    db.add(user)
    db.commit()
    return {"ok": True}


@app.post("/documents/upload", response_model=DocumentOut)
def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(""),
    language: str = Form("az"),
    db: Session = Depends(get_db),
):
    # QEYD: bu handler qəsdən "async def" DEYİL. İçəridə Ollama-ya (analiz + hər chunk üçün
    # embedding) uzun sürən, tam SİNXRON/bloklayıcı urllib çağırışları var. "async def" olsaydı,
    # FastAPI bunu birbaşa tək asyncio event loop-unda işlədərdi — bu bloklayıcı çağırış davam
    # etdiyi müddətdə (indi hər biri 300 saniyəyə qədər çəkə bilər) server HEÇ BİR başqa sorğunu
    # (başqa istifadəçinin sualını, hətta sadə səhifə yüklənməsini) qəbul edə bilməzdi — məhz
    # "server tamam donub" simptomunun səbəbi bu idi. Sinxron "def" olduqda isə FastAPI/Starlette
    # bunu avtomatik ayrıca worker thread-də işlədir, event loop sərbəst qalır.
    original_name = file.filename or "document"
    suffix = Path(original_name).suffix.lower()
    if suffix not in {".pdf", ".docx", ".txt", ".jpg", ".jpeg", ".png"}:
        raise HTTPException(400, "PDF, DOCX, TXT və ya şəkil (JPG/PNG) yükləyin.")
    path = Path(settings.upload_dir) / f"{uuid.uuid4()}{suffix}"
    with path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        text = extract_text(path)
    except HTTPException:
        path.unlink(missing_ok=True)
        raise
    except Exception:
        path.unlink(missing_ok=True)
        raise HTTPException(400, "Sənəd oxuna bilmədi. Fayl zədələnmiş və ya dəstəklənməyən formatda ola bilər.")

    if not text.strip():
        path.unlink(missing_ok=True)
        raise HTTPException(400, "Sənəddən oxuna bilən mətn çıxmadı (OCR də nəticə vermədi). Şəkil keyfiyyəti aşağı ola bilər.")

    try:
        result = analyze_document(text, language=language)
    except Exception:
        path.unlink(missing_ok=True)
        raise HTTPException(500, "Sənəd analiz edilərkən xəta baş verdi. Bir azdan yenidən cəhd edin.")

    client_name = result["client_name"].strip() or "Naməlum müştəri"
    if len(client_name) > 500:  # DB sütununun ölçüsü (clients.full_name) — əlavə təhlükəsizlik.
        client_name = client_name[:497].rstrip() + "…"
    client = db.scalar(select(Client).where(Client.full_name == client_name))
    if not client:
        client = Client(full_name=client_name)
        db.add(client)
        db.flush()
    document = Document(
        original_filename=original_name,
        stored_path=str(path),
        document_type=document_type or result["document_type"],
        extracted_text=text,
        client_id=client.id,
        embedding=embed_document_text(text),
    )
    db.add(document)
    db.flush()
    document.analysis = Analysis(
        document_id=document.id,
        summary=result["summary"],
        risks=json.dumps(result["risks"], ensure_ascii=False),
        grammar_issues=json.dumps(result["grammar_issues"], ensure_ascii=False),
        extracted_dates=json.dumps(result["extracted_dates"], ensure_ascii=False),
        risk_score=result["risk_score"],
    )

    # Dəqiq RAG axtarışı üçün sənədi LlamaIndex ilə kiçik parçalara bölüb hər birini
    # ayrıca embed edirik (bütöv-sənəd embedding-i yuxarıda /search üçün saxlanılır).
    for idx, chunk in enumerate(chunk_text(text)):
        db.add(DocumentChunk(
            document_id=document.id,
            chunk_index=idx,
            text=chunk,
            embedding=embed_document_text(chunk),
        ))

    contract_end_date = (result.get("contract_end_date") or "").strip()
    if contract_end_date:
        try:
            parsed_date = datetime.strptime(contract_end_date, "%Y-%m-%d")
        except ValueError:
            parsed_date = None
        if parsed_date:
            db.add(CourtEvent(
                client_id=client.id,
                court_date=parsed_date,
                note=f"Müqavilə bitmə tarixi — {original_name}",
                event_type="contract",
            ))

    db.commit()
    db.refresh(document)
    return serialize(document)


@app.get("/search", response_model=list[DocumentOut])
def search_client(q: str, db: Session = Depends(get_db)):
    # Hərfi axtarış — müştəri adı/fayl adı/mətn üzərində birbaşa alt-sətir uyğunluğu.
    # Ad və fayl adları üçün bu, ən etibarlı üsuldur.
    like = f"%{q}%"
    literal_docs = list(db.scalars(
        select(Document)
        .join(Document.client)
        .options(joinedload(Document.client), joinedload(Document.analysis))
        .where(or_(
            Client.full_name.ilike(like),
            Document.extracted_text.ilike(like),
            Document.original_filename.ilike(like),
        ))
        .order_by(Document.created_at.desc())
    ).unique().all())

    # Semantik axtarış — sorğunun MƏNASINA yaxın sənəd mətnlərini tapır (embedding oxşarlığı).
    # Hərfi uyğunluq tapılmasa belə, məzmunca əlaqəli sənədləri gətirir.
    semantic_docs: list[Document] = []
    query_embedding = embed_query(q)
    if query_embedding is not None:
        distance = Document.embedding.cosine_distance(query_embedding)
        semantic_docs = list(db.scalars(
            select(Document)
            .options(joinedload(Document.client), joinedload(Document.analysis))
            .where(Document.embedding.is_not(None))
            .where(distance < 0.6)  # 0=eyni məna, 2=tam əlaqəsiz; 0.6 orta-sərt astanadır
            .order_by(distance)
            .limit(10)
        ).unique().all())

    seen_ids: set[int] = set()
    merged: list[Document] = []
    for doc in literal_docs + semantic_docs:
        if doc.id not in seen_ids:
            seen_ids.add(doc.id)
            merged.append(doc)
    return [serialize(document) for document in merged]


@app.post("/admin/backfill-embeddings")
def backfill_embeddings(force: bool = False, db: Session = Depends(get_db)):
    """Köhnə (pgvector/chunking əlavə olunmazdan əvvəl yüklənmiş) sənədləri sinxronlaşdırır:
    - Document.embedding: force=false olduqda yalnız boş olanları doldurur; force=true
      olduqda hamısını "search_document:" prefiksi ilə YENİDƏN hesablayır (embedding
      generasiya üsulu dəyişəndə köhnə sənədləri yeniləmək üçün).
    - DocumentChunk: hələ heç bir parçası olmayan sənədləri LlamaIndex SentenceSplitter
      ilə bölüb hər parçanı embed edir (force=true olduqda mövcud parçalar silinib yenidən yaradılır).
    """
    doc_query = select(Document) if force else select(Document).where(Document.embedding.is_(None))
    docs_to_embed = db.scalars(doc_query).all()
    doc_updated = 0
    for doc in docs_to_embed:
        embedding = embed_document_text(doc.extracted_text)
        if embedding is not None:
            doc.embedding = embedding
            doc_updated += 1

    all_docs = db.scalars(select(Document)).all()
    chunks_created = 0
    for doc in all_docs:
        has_chunks = db.scalar(select(DocumentChunk.id).where(DocumentChunk.document_id == doc.id).limit(1))
        if has_chunks and not force:
            continue
        if has_chunks and force:
            db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == doc.id))
        for idx, chunk in enumerate(chunk_text(doc.extracted_text)):
            db.add(DocumentChunk(
                document_id=doc.id,
                chunk_index=idx,
                text=chunk,
                embedding=embed_document_text(chunk),
            ))
            chunks_created += 1

    db.commit()
    return {
        "documents_embedded": doc_updated,
        "chunks_created": chunks_created,
        "total_documents": len(all_docs),
    }


@app.get("/documents", response_model=list[DocumentOut])
def list_documents(limit: int = 50, db: Session = Depends(get_db)):
    docs = db.scalars(
        select(Document)
        .options(joinedload(Document.client), joinedload(Document.analysis))
        .order_by(Document.created_at.desc())
        .limit(limit)
    ).unique().all()
    return [serialize(document) for document in docs]


RAG_RETRIEVE_K = 10  # pgvector ilə çəkilən ilkin namizəd (chunk) sayı
RAG_FINAL_K = 5  # LLMRerank-dan sonra qalan, cavab üçün istifadə olunan chunk sayı

# Salamlaşma/ümumi söhbət — bunlar üçün vector DB-yə heç getmirik, birbaşa cavab veririk.
GREETING_PATTERNS = [
    "salam", "necəsən", "necesen", "sağol", "sagol", "sağ ol", "sag ol",
    "hi", "hello", "hey", "good morning", "good evening",
    "təşəkkür", "tesekkur", "sağ olun", "sag olun", "xudahafiz", "sabahın xeyir",
    "axşamın xeyir", "günortanız xeyir", "здравствуй", "привет", "спасибо",
]


def _looks_like_greeting(question: str) -> bool:
    """Qısa (≤4 sözlük) salamlaşma/nəzakət ifadələrini aşkarlayır — real sual bunlardan
    daha uzun olur, ona görə uzun mesajlar (məs. 'salam, bu sənəddə risk varmı?') bura düşmür."""
    normalized = question.lower().strip().strip("?!.,")
    if len(normalized.split()) > 4:
        return False
    return any(pattern in normalized for pattern in GREETING_PATTERNS)


GREETING_ANSWERS = {
    "az": "Salam! Mən LegalAI köməkçisiyəm. Sənəd yükləyə, onun risklərini/xülasəsini soruşa və ya müddətləri yoxlaya bilərsiniz.",
    "en": "Hello! I'm the LegalAI assistant. You can upload a document, ask about its risks or summary, or check deadlines.",
    "ru": "Здравствуйте! Я ассистент LegalAI. Вы можете загрузить документ, спросить о рисках/резюме или проверить сроки.",
}

AZ_VOWELS = set("aəeıioöuüAƏEIİOÖUÜ")
# Tək sözlük real sual/əmr ola bilən qısa sözlər — bunlar "mənasız" sayılmır.
KNOWN_SHORT_WORDS = {
    "kim", "nə", "ne", "niyə", "niye", "harda", "harada", "nə vaxt",
    "necə", "nece", "hansı", "hansi", "risk", "risklər", "risqler",
    "xülasə", "xulase", "salam", "hi", "hello", "summary", "kömək", "komek",
    "tarix", "vaxt", "sənəd", "senedi", "müqavilə", "muqavile",
}


def _looks_like_gibberish(question: str) -> bool:
    """Sadə yoxlama: (1) sait daşımayan hərf yığını (məs. 'mghbg'), (2) tanınmayan tək söz — mənasız sayılır."""
    text = question.strip()
    letters = [c for c in text if c.isalpha()]
    if len(letters) < 3:
        return True
    if not any(c in AZ_VOWELS for c in letters):
        return True
    words = text.split()
    if len(words) == 1 and text.lower().strip("?!.,") not in KNOWN_SHORT_WORDS:
        return True
    return False


# "Yazı xətası yoxla / xülasə ver / risk var?" kimi suallar üçün LLM-ə yenidən sormaq
# əvəzinə, sənəd yükləndikdə DƏQİQ bunun üçün yazılmış prompt-la (bax: ai.py SYSTEM_PROMPT)
# artıq hesablanmış nəticələri (Analysis) birbaşa göstəririk. Bu, iki səbəbdən daha
# etibarlıdır: (1) analiz zamanı istifadə olunan prompt konkret bu tapşırıq üçün yazılıb,
# çat zamanı istifadə olunan ümumi sual-cavab prompt-undan güclüdür; (2) LLM-i təkrar
# çağırmadığı üçün nə model zəifliyindən, nə də uzun sənədin kontekst pəncərəsinə
# sığmamasından əziyyət çəkir.
GRAMMAR_KEYWORDS = [
    "yazı xəta", "yazi xeta", "qrammatika", "qramatika", "orfoqrafiya",
    "syntax", "grammar", "spelling", "орфограф", "грамматик", "синтаксис",
]
SUMMARY_KEYWORDS = [
    "xülasə", "xulase", "nə haqqında", "ne haqqinda", "nə var", "ne var",
    "nə baş verib", "ne bas verib", "nə barədədir", "ne barededir",
    "summary", "summarize", "what is this", "what happened", "what's in",
    "о чём", "что произошло", "краткое содержание", "о чем",
]
RISK_KEYWORDS = ["risk", "риск"]


def _looks_like_analysis_question(question: str) -> str | None:
    """Sual grammar/summary/risk niyyətlərindən hansına uyğundursa onu qaytarır, yoxdursa None."""
    normalized = question.lower()
    if any(k in normalized for k in GRAMMAR_KEYWORDS):
        return "grammar"
    if any(k in normalized for k in RISK_KEYWORDS):
        return "risk"
    if any(k in normalized for k in SUMMARY_KEYWORDS):
        return "summary"
    return None


_ANALYSIS_LABELS = {
    "az": {"summary": "Xülasə", "risk_score": "Risk balı", "risks": "Risklər", "none": "Aşkar edilmədi.", "grammar": "Yazı/qrammatika xətaları"},
    "en": {"summary": "Summary", "risk_score": "Risk score", "risks": "Risks", "none": "None found.", "grammar": "Grammar/syntax issues"},
    "ru": {"summary": "Резюме", "risk_score": "Оценка риска", "risks": "Риски", "none": "Не обнаружено.", "grammar": "Орфографические/грамматические ошибки"},
}


def _direct_analysis_answer(intent: str, documents: list[Document], language: str) -> str | None:
    labels = _ANALYSIS_LABELS.get(language, _ANALYSIS_LABELS["az"])
    blocks = []
    for doc in documents:
        if not doc.analysis:
            continue
        lines = [f"📄 {doc.original_filename}"]
        if intent == "grammar":
            issues = json.loads(doc.analysis.grammar_issues) if doc.analysis.grammar_issues else []
            lines.append(f"{labels['grammar']}: " + ("; ".join(issues) if issues else labels["none"]))
        elif intent == "risk":
            risks = json.loads(doc.analysis.risks) if doc.analysis.risks else []
            lines.append(f"{labels['risk_score']}: {doc.analysis.risk_score}%")
            lines.append(f"{labels['risks']}: " + ("; ".join(risks) if risks else labels["none"]))
        else:  # summary
            lines.append(f"{labels['summary']}: {doc.analysis.summary}")
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks) if blocks else None


def _document_block(doc: Document) -> str:
    parts = [f"Sənəd: {doc.original_filename}"]
    if doc.analysis:
        risks = json.loads(doc.analysis.risks) if doc.analysis.risks else []
        parts.append(f"AI-nin əvvəlcədən müəyyənləşdirdiyi risk balı: {doc.analysis.risk_score}%")
        parts.append("AI-nin əvvəlcədən müəyyənləşdirdiyi konkret risklər: " + ("; ".join(risks) if risks else "tapılmayıb"))
        parts.append(f"Xülasə: {doc.analysis.summary}")
    parts.append(f"Tam mətn:\n{doc.extracted_text}")
    return "\n".join(parts)


def _sse(stage: str, **fields) -> str:
    return json.dumps({"stage": stage, **fields}, ensure_ascii=False) + "\n"


def _ask_stream(payload: AskRequest):
    """Backend RAG addımlarını NDJSON axını kimi göndərir ki, UI hansı mərhələdə
    olduğumuzu göstərə bilsin (embedding, axtarış, rerank, cavab yaradılması).
    DB sessiyasını burda özümüz açıb bağlayırıq — StreamingResponse Depends(get_db)
    sessiyasını streaming başlamazdan əvvəl bağlaya bilər, ona görə FastAPI-nin
    inyeksiyasına etibar etmirik."""
    db = SessionLocal()
    try:
        if _looks_like_gibberish(payload.question):
            yield _sse("done", answer="Sualınız aydın olmadı. Zəhmət olmasa sualınızı daha aydın yazın.", sources=[])
            return

        if _looks_like_greeting(payload.question):
            yield _sse("done", answer=GREETING_ANSWERS.get(payload.language, GREETING_ANSWERS["az"]), sources=[])
            return

        yield _sse("progress", message="Sual analiz edilir...")

        normalized_question = payload.question.lower()
        if ("məhkəmə" in normalized_question or "mehkeme" in normalized_question) and ("gün" in normalized_question or "tarix" in normalized_question):
            yield _sse("progress", message="Məhkəmə günləri yoxlanılır...")
            events = db.scalars(
                select(CourtEvent)
                .where(CourtEvent.court_date >= datetime.now(), CourtEvent.event_type == "court")
                .order_by(CourtEvent.court_date)
            ).all()
            if not events:
                yield _sse("done", answer="Yaxınlaşan məhkəmə günü qeyd olunmayıb.", sources=[])
                return
            event_text = "; ".join(f"{event.court_date.strftime('%d.%m.%Y %H:%M')} — iş № {event.case_number or 'qeyd olunmayıb'}" for event in events)
            yield _sse("done", answer=f"Yaxınlaşan məhkəmə günləri: {event_text}", sources=[])
            return

        client_filter = payload.client_name.strip()
        candidates: list[tuple[str, float, dict]] = []

        if client_filter:
            yield _sse("progress", message=f"'{client_filter}' üzrə sənədlər axtarılır...")
            documents = list(db.scalars(
                select(Document)
                .options(joinedload(Document.client), joinedload(Document.analysis))
                .join(Document.client)
                .where(Client.full_name.ilike(f"%{client_filter}%"))
            ).unique().all())

            # Yazı xətası / xülasə / risk sualları — LLM-i yenidən çağırmadan, yükləmə
            # zamanı artıq hesablanmış (və daha etibarlı) Analysis nəticələrindən cavab veririk.
            intent = _looks_like_analysis_question(payload.question)
            if intent and documents:
                direct_answer = _direct_analysis_answer(intent, documents, payload.language)
                if direct_answer:
                    yield _sse("done", answer=direct_answer, sources=[{"filename": d.original_filename, "score": 100.0} for d in documents])
                    return

            candidates = [(_document_block(doc), 1.0, {"filename": doc.original_filename}) for doc in documents]

        if not candidates:
            yield _sse("progress", message="Sualın mənası embedding-ə çevrilir (LlamaIndex + Ollama)...")
            query_embedding = embed_query(payload.question)
            if query_embedding is not None:
                yield _sse("progress", message="Uyğun sənəd parçaları pgvector ilə axtarılır...")
                distance = DocumentChunk.embedding.cosine_distance(query_embedding)
                stmt = (
                    select(DocumentChunk, distance.label("distance"))
                    .options(joinedload(DocumentChunk.document).joinedload(Document.analysis))
                    .where(DocumentChunk.embedding.is_not(None))
                    .order_by(distance)
                    .limit(RAG_RETRIEVE_K)
                )
                rows = db.execute(stmt).unique().all()
                # cosine_distance: 0 = eyni məna, 2 = tam əlaqəsiz → relevanslıq balına çeviririk.
                for chunk, dist in rows:
                    score = max(0.0, 1.0 - dist)
                    header = f"Sənəd: {chunk.document.original_filename}"
                    if chunk.document.analysis:
                        header += f" | Risk balı: {chunk.document.analysis.risk_score}%"
                    candidates.append((f"{header}\nMətn parçası:\n{chunk.text}", score, {"filename": chunk.document.original_filename}))

            if not candidates:
                yield _sse("progress", message="Uyğun parça tapılmadı, son yüklənmiş sənədlərə keçilir...")
                fallback_docs = list(db.scalars(
                    select(Document)
                    .options(joinedload(Document.client), joinedload(Document.analysis))
                    .order_by(Document.created_at.desc())
                    .limit(RAG_FINAL_K)
                ).unique().all())
                candidates = [(_document_block(doc), 0.5, {"filename": doc.original_filename}) for doc in fallback_docs]

        if not candidates:
            yield _sse("done", answer="Bu sualı cavablandırmaq üçün əvvəlcə sənəd yükləyin və ya düzgün müştəri adı yazın.", sources=[])
            return

        if len(candidates) > RAG_FINAL_K:
            yield _sse("progress", message=f"{len(candidates)} nəticə LLMRerank ilə ən yaxşı {RAG_FINAL_K}-ə endirilir...")

        yield _sse("progress", message="Cavab LLM ilə yaradılır (LlamaIndex response synthesizer)...")
        try:
            answer, sources = rerank_and_answer(payload.question, candidates, language=payload.language, top_n=RAG_FINAL_K)
        except Exception:
            logger.exception("LlamaIndex RAG xətası, ənənəvi Ollama yoluna keçirik")
            yield _sse("progress", message="LlamaIndex əlçatan olmadı, ehtiyat yol (birbaşa Ollama) sınanılır...")
            context = "\n\n---\n\n".join(text for text, _, _ in candidates[:RAG_FINAL_K])
            answer = answer_question(payload.question, context, language=payload.language)
            seen_filenames: set[str] = set()
            sources = []
            for text, score, meta in sorted(candidates, key=lambda c: c[1], reverse=True):
                filename = meta.get("filename", "sənəd")
                if filename in seen_filenames:
                    continue
                seen_filenames.add(filename)
                sources.append({"filename": filename, "score": round(score * 100, 1)})
                if len(sources) >= RAG_FINAL_K:
                    break

        yield _sse("done", answer=answer, sources=sources)
    finally:
        db.close()


@app.post("/ask")
def ask_question(payload: AskRequest):
    return StreamingResponse(_ask_stream(payload), media_type="application/x-ndjson")


@app.post("/court-events")
def add_court_event(payload: CourtEventCreate, db: Session = Depends(get_db)):
    client = db.scalar(select(Client).where(Client.full_name == payload.client_name))
    if not client:
        client = Client(full_name=payload.client_name)
        db.add(client)
        db.flush()
    event = CourtEvent(
        client_id=client.id,
        court_date=payload.court_date,
        case_number=payload.case_number,
        note=payload.note,
        event_type=payload.event_type or "court",
    )
    db.add(event)
    db.commit()
    return {"id": event.id, "message": "Tarix əlavə edildi"}


@app.get("/court-events/upcoming", response_model=list[DeadlineOut])
def upcoming_court_events(client_name: str = "", db: Session = Depends(get_db)):
    statement = (
        select(CourtEvent)
        .join(CourtEvent.client)
        .options(joinedload(CourtEvent.client))
        .where(CourtEvent.court_date >= datetime.now())
        .order_by(CourtEvent.court_date)
    )
    if client_name.strip():
        statement = statement.where(Client.full_name.ilike(f"%{client_name.strip()}%"))
    events = db.scalars(statement).unique().all()
    return [
        {
            "id": e.id,
            "client_name": e.client.full_name,
            "court_date": e.court_date,
            "case_number": e.case_number,
            "note": e.note,
            "event_type": e.event_type,
        }
        for e in events
    ]


# ---- n8n Telegram bildirişləri üçün "alert" endpointləri ----
# n8n bunları müəyyən aralıqla (Schedule Trigger) yoxlayır, yeni tapılanı Telegram-a göndərir,
# sonra "mark-notified" ilə işarələyir ki, eyni bildiriş təkrar getməsin.

@app.get("/alerts/risky-documents")
def risky_documents(threshold: float = 70, db: Session = Depends(get_db)):
    docs = db.scalars(
        select(Document)
        .join(Document.analysis)
        .options(joinedload(Document.client), joinedload(Document.analysis))
        .where(Analysis.risk_score >= threshold, Analysis.notified.is_(False))
        .order_by(Document.created_at.desc())
    ).unique().all()
    return [
        {
            "document_id": d.id,
            "client_name": d.client.full_name,
            "original_filename": d.original_filename,
            "risk_score": d.analysis.risk_score,
            "summary": d.analysis.summary,
        }
        for d in docs
    ]


@app.post("/alerts/risky-documents/{document_id}/mark-notified")
def mark_risky_document_notified(document_id: int, db: Session = Depends(get_db)):
    analysis = db.scalar(select(Analysis).where(Analysis.document_id == document_id))
    if not analysis:
        raise HTTPException(404, "Sənəd tapılmadı.")
    analysis.notified = True
    db.commit()
    return {"message": "Bildiriş qeyd olundu"}


@app.get("/alerts/deadlines")
def deadline_alerts(days: int = 3, db: Session = Depends(get_db)):
    cutoff = datetime.now() + timedelta(days=days)
    events = db.scalars(
        select(CourtEvent)
        .join(CourtEvent.client)
        .options(joinedload(CourtEvent.client))
        .where(
            CourtEvent.court_date >= datetime.now(),
            CourtEvent.court_date <= cutoff,
            CourtEvent.notified.is_(False),
        )
        .order_by(CourtEvent.court_date)
    ).unique().all()
    return [
        {
            "event_id": e.id,
            "client_name": e.client.full_name,
            "court_date": e.court_date,
            "case_number": e.case_number,
            "note": e.note,
            "event_type": e.event_type,
        }
        for e in events
    ]


@app.post("/alerts/deadlines/{event_id}/mark-notified")
def mark_deadline_notified(event_id: int, db: Session = Depends(get_db)):
    event = db.get(CourtEvent, event_id)
    if not event:
        raise HTTPException(404, "Tarix tapılmadı.")
    event.notified = True
    db.commit()
    return {"message": "Bildiriş qeyd olundu"}


@app.get("/alerts/keyword-documents")
def keyword_documents(phrase: str, db: Session = Depends(get_db)):
    """Sənəd mətnində konkret ifadə (məs. 'azadlıqdan məhrum etmə cəzası təyin edilmişdir') keçirsə tapır."""
    docs = db.scalars(
        select(Document)
        .join(Document.client)
        .options(joinedload(Document.client))
        .where(Document.extracted_text.ilike(f"%{phrase.strip()}%"), Document.keyword_notified.is_(False))
        .order_by(Document.created_at.desc())
    ).unique().all()
    return [
        {
            "document_id": d.id,
            "client_name": d.client.full_name,
            "original_filename": d.original_filename,
            "matched_phrase": phrase,
        }
        for d in docs
    ]


@app.post("/alerts/keyword-documents/{document_id}/mark-notified")
def mark_keyword_document_notified(document_id: int, db: Session = Depends(get_db)):
    document = db.get(Document, document_id)
    if not document:
        raise HTTPException(404, "Sənəd tapılmadı.")
    document.keyword_notified = True
    db.commit()
    return {"message": "Bildiriş qeyd olundu"}
